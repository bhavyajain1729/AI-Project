import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown

def add_heading_style(doc, text, level=1):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    heading.style = f'Heading {level}'
    return heading

def add_table_of_contents(doc):
    """Add a table of contents"""
    doc.add_page_break()
    heading = doc.add_heading('Table of Contents', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add TOC paragraph
    toc = doc.add_paragraph()
    run = toc.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "TOC \\o '1-2' \\h \\z \\u"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    doc.add_page_break()

def parse_markdown_to_dict(file_path):
    """Parse markdown file and extract structured data"""
    data = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                data['title'] = line.replace('# ', '').strip()
            elif line.startswith('## '):
                current_section = line.replace('## ', '').strip()
                data[current_section] = []
            elif line.startswith('**') and ':' in line:
                key_val = line.replace('**', '').split(':', 1)
                if len(key_val) == 2:
                    data[key_val[0].strip()] = key_val[1].strip()
            elif line.startswith('- '):
                if current_section:
                    data[current_section].append(line.replace('- ', '').strip())
    except Exception as e:
        print(f"Error parsing {file_path}: {e}")
    
    return data

def add_module_content(doc, file_path, category):
    """Add module content to document"""
    parsed = parse_markdown_to_dict(file_path)
    
    if not parsed:
        return
    
    # Add module title
    title = parsed.get('title', os.path.basename(file_path))
    doc.add_heading(title, level=2)
    
    # Add source file
    source = parsed.get('Source file', '')
    if source:
        doc.add_paragraph(f"Source: {source}", style='Normal')
    
    # Add summary
    if 'Summary' in parsed:
        doc.add_heading('Summary', level=3)
        doc.add_paragraph(parsed['Summary'])
    
    # Add module info
    if 'Module' in parsed:
        doc.add_heading('Module', level=3)
        doc.add_paragraph(parsed['Module'])
    
    # Add procedures
    if 'Procedures' in parsed and parsed['Procedures']:
        doc.add_heading('Procedures', level=3)
        for proc in parsed['Procedures']:
            doc.add_paragraph(proc, style='List Bullet')
    
    # Add tables
    if 'Tables' in parsed and parsed['Tables']:
        doc.add_heading('Tables', level=3)
        for table in parsed['Tables']:
            doc.add_paragraph(table, style='List Bullet')
    
    # Add entities
    if 'Entities' in parsed and parsed['Entities']:
        doc.add_heading('Entities', level=3)
        for entity in parsed['Entities']:
            doc.add_paragraph(entity, style='List Bullet')
    
    # Add dependencies
    if 'Dependencies' in parsed and parsed['Dependencies']:
        doc.add_heading('Dependencies', level=3)
        for dep in parsed['Dependencies']:
            doc.add_paragraph(dep, style='List Bullet')
    
    doc.add_paragraph()  # Add space

def create_word_documentation():
    """Create Word document from all markdown files"""
    doc = Document()
    
    # Add title page
    title = doc.add_heading('Project Modules Documentation', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Module Reference')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add generation date
    from datetime import datetime
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add TOC
    add_table_of_contents(doc)
    
    # Define modules and their paths
    modules = {
        'DOA': r'c:\Users\bhavya.jain\Desktop\AI_Project\output\docs\DOA',
        'Inventory': r'c:\Users\bhavya.jain\Desktop\AI_Project\output\docs\Inventory',
        'Reporting': r'c:\Users\bhavya.jain\Desktop\AI_Project\output\docs\Reporting',
        'Utilities': r'c:\Users\bhavya.jain\Desktop\AI_Project\output\docs\Utilities'
    }
    
    # Process each module category
    for module_name, module_path in modules.items():
        doc.add_heading(module_name, level=1)
        
        # Get all markdown files (exclude _1 versions)
        files = [f for f in os.listdir(module_path) if f.endswith('.md') and not f.endswith('_1.md')]
        files.sort()
        
        for file_name in files:
            file_path = os.path.join(module_path, file_name)
            add_module_content(doc, file_path, module_name)
        
        doc.add_page_break()
    
    # Save document
    output_path = r'c:\Users\bhavya.jain\Desktop\AI_Project\Module_Documentation.docx'
    doc.save(output_path)
    print(f"✓ Documentation generated successfully!")
    print(f"✓ Saved to: {output_path}")

if __name__ == '__main__':
    create_word_documentation()
